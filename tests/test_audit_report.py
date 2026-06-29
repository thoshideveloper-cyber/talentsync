"""
Workstream 1 — Audit Report endpoint tests.

GET /api/records/{id}/audit-report
  - Any authenticated role may download (not just approver/admin)
  - ?format=docx returns a DOCX binary
  - ?format=pdf returns 422 (not yet supported)
  - 404 for unknown record
  - 401 for unauthenticated request
  - DOCX content contains required sections
  - Audit log row written after generation
"""
import uuid as _uuid

import pytest

_HIGH_RISK_JD = (
    "Software Engineer. Age below 30 only. Male candidates preferred. "
    "5+ years Python, AWS. Salary: 18-25 LPA. Bengaluru."
)

_CLEAN_JD = (
    "Senior Data Analyst. 4+ years SQL, Python, Tableau. "
    "Salary: 15-22 LPA. Mumbai (Hybrid). Equal opportunity employer."
)


async def _upload(client, text: str) -> str:
    resp = await client.post("/api/extract", json={"text": text})
    assert resp.status_code == 200, resp.text
    return resp.json()["id"]


# ── Core happy-path tests ─────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_audit_report_generates_for_high_risk_jd(recruiter_client):
    """Recruiter (any role) can download. Response is DOCX bytes."""
    record_id = await _upload(recruiter_client, _HIGH_RISK_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert len(resp.content) > 1000  # non-trivial DOCX


@pytest.mark.asyncio
async def test_audit_report_generates_for_clean_jd(recruiter_client):
    record_id = await _upload(recruiter_client, _CLEAN_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report")
    assert resp.status_code == 200
    assert len(resp.content) > 500


@pytest.mark.asyncio
async def test_audit_report_contains_required_sections(recruiter_client):
    """Parse the DOCX and check all 7 required sections are present."""
    import io
    from docx import Document

    record_id = await _upload(recruiter_client, _HIGH_RISK_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report")
    assert resp.status_code == 200

    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 1. Cover
    assert "TalentSync Compliance Audit Report" in full_text
    # 2. Verdict summary / Executive Summary
    assert "Executive Summary" in full_text
    # 3. Findings section
    assert "Compliance Findings" in full_text
    # 4. Methodology disclosure
    assert "Methodology" in full_text
    # 5. Risk-gradient key
    assert "Risk-Gradient Key" in full_text
    # 6. Audit trail
    assert "Audit Trail" in full_text
    # 7. Footer disclaimer
    assert "not a legal opinion" in full_text.lower() or "Not a legal opinion" in full_text


@pytest.mark.asyncio
async def test_audit_report_clean_jd_no_findings_message(recruiter_client):
    """A clean JD should produce 'No compliance issues detected' in the report."""
    import io
    from docx import Document

    record_id = await _upload(recruiter_client, _CLEAN_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report")
    assert resp.status_code == 200

    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "No compliance issues detected" in full_text


@pytest.mark.asyncio
async def test_audit_report_methodology_text_present(recruiter_client):
    """The methodology disclosure must include the recall caveat text."""
    import io
    from docx import Document

    record_id = await _upload(recruiter_client, _HIGH_RISK_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report")
    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Recall" in full_text


# ── Format validation ─────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_audit_report_unsupported_format_returns_422(recruiter_client):
    record_id = await _upload(recruiter_client, _CLEAN_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report?format=pdf")
    assert resp.status_code == 422


@pytest.mark.asyncio
async def test_audit_report_format_docx_explicit(recruiter_client):
    record_id = await _upload(recruiter_client, _CLEAN_JD)
    resp = await recruiter_client.get(f"/api/records/{record_id}/audit-report?format=docx")
    assert resp.status_code == 200


# ── Error cases ───────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_audit_report_unknown_record_returns_404(recruiter_client):
    resp = await recruiter_client.get(f"/api/records/{_uuid.uuid4()}/audit-report")
    assert resp.status_code == 404


@pytest.mark.asyncio
async def test_audit_report_requires_auth(client):
    resp = await client.get(f"/api/records/{_uuid.uuid4()}/audit-report")
    assert resp.status_code == 401


# ── Audit log written ─────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_audit_report_writes_audit_log(recruiter_client, admin_client):
    """Downloading the report must write an audit_report.generated row."""
    record_id = await _upload(recruiter_client, _HIGH_RISK_JD)
    await recruiter_client.get(f"/api/records/{record_id}/audit-report")

    # Audit log is visible to admin
    audit_resp = await admin_client.get(f"/api/records/{record_id}/audit")
    assert audit_resp.status_code == 200
    actions = [row["action"] for row in audit_resp.json()]
    assert "audit_report.generated" in actions
